"""知识库文档处理服务。"""

import json
import uuid
from pathlib import Path

from docx import Document as DocxDocument
from flask import current_app
from pypdf import PdfReader
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename

from app.extensions import db
from app.models import KbChunk, KbDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter


class DocumentService:
    """负责文件保存、文本抽取与切片入库。"""

    ALLOWED_EXTENSIONS = {"txt", "md", "pdf", "docx"}

    @classmethod
    def create_document(cls, upload_file: FileStorage, title: str, knowledge_base_id: int, user_id: int):
        """保存上传文件并创建文档记录。"""

        file_name = secure_filename(upload_file.filename or "")
        if not file_name:
            raise ValueError("请选择需要上传的文件")

        extension = file_name.rsplit(".", 1)[-1].lower()
        if extension not in cls.ALLOWED_EXTENSIONS:
            raise ValueError("当前仅支持 txt、md、pdf、docx 格式文件")

        save_name = f"{uuid.uuid4().hex}.{extension}"
        save_path = Path(current_app.config["UPLOAD_DIR"]) / save_name
        upload_file.save(save_path)

        document = KbDocument(
            knowledge_base_id=knowledge_base_id,
            title=title or Path(file_name).stem,
            original_name=file_name,
            file_type=extension,
            file_path=str(save_path),
            process_status="pending",
            created_by=user_id,
        )
        db.session.add(document)
        db.session.commit()
        return document

    @classmethod
    def extract_text(cls, file_path: str, file_type: str) -> str:
        """根据文件类型抽取纯文本内容。"""

        resolved_path = Path(file_path)
        if not resolved_path.is_absolute() and not resolved_path.exists():
            candidate_paths = [
                Path(current_app.root_path).parent / file_path,
                Path(current_app.root_path).parent.parent / file_path,
            ]
            for candidate in candidate_paths:
                if candidate.exists():
                    resolved_path = candidate
                    break

        if not resolved_path.exists():
            raise ValueError(f"文件不存在：{file_path}")

        if file_type in {"txt", "md"}:
            return resolved_path.read_text(encoding="utf-8")

        if file_type == "pdf":
            reader = PdfReader(str(resolved_path))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n".join(pages).strip()

        if file_type == "docx":
            doc = DocxDocument(str(resolved_path))
            return "\n".join([paragraph.text for paragraph in doc.paragraphs]).strip()

        raise ValueError("不支持的文件类型")

    @classmethod
    def split_and_store_chunks(cls, document: KbDocument):
        """将文档切片后写入数据库，供向量索引服务进一步处理。"""

        content = cls.extract_text(document.file_path, document.file_type).strip()
        if not content:
            raise ValueError("文档内容为空，无法建立知识库索引")

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=100,
            separators=["\n\n", "\n", "。", "！", "？", "；", "，", " "],
        )
        chunks = splitter.split_text(content)

        KbChunk.query.filter_by(document_id=document.id).delete()
        db.session.flush()

        chunk_records = []
        for index, chunk_text in enumerate(chunks, start=1):
            meta_json = json.dumps(
                {
                    "documentTitle": document.title,
                    "knowledgeBaseId": document.knowledge_base_id,
                    "sequence": index,
                },
                ensure_ascii=False,
            )
            chunk_record = KbChunk(
                document_id=document.id,
                sequence=index,
                vector_id=f"doc-{document.id}-chunk-{index}",
                content=chunk_text,
                char_count=len(chunk_text),
                meta_json=meta_json,
            )
            chunk_records.append(chunk_record)

        db.session.add_all(chunk_records)
        document.chunk_count = len(chunk_records)
        document.process_status = "processed"
        db.session.commit()
        return chunk_records

    @staticmethod
    def mark_failed(document: KbDocument):
        """当文档处理失败时记录状态。"""

        document.process_status = "failed"
        db.session.commit()
